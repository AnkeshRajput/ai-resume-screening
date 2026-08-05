"""
Resume parsing service for extracting text from PDF and DOCX documents.
"""

import io
import os
import re
import docx
import pdfplumber


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF file using pdfplumber across all pages.

    Args:
        file_bytes (bytes): The raw file bytes of the PDF.

    Returns:
        str: Extracted combined text content.

    Raises:
        ValueError: If PDF is encrypted, corrupted, or unreadable.
    """
    extracted_pages = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            if not pdf.pages:
                raise ValueError("PDF file contains no pages.")
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    extracted_pages.append(page_text)
    except Exception as e:
        error_msg = str(e).lower()
        if "password" in error_msg or "encrypted" in error_msg:
            raise ValueError("The uploaded PDF is password-protected. Please upload an unprotected PDF.")
        raise ValueError(f"Failed to read PDF file: {str(e)}")

    return "\n".join(extracted_pages)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX file using python-docx.

    Args:
        file_bytes (bytes): The raw file bytes of the DOCX document.

    Returns:
        str: Extracted combined text content.

    Raises:
        ValueError: If DOCX document is corrupted or unreadable.
    """
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        
        # Also extract text from tables inside docx if present
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    table_texts.append(" | ".join(row_text))

        combined_text = "\n".join(paragraphs + table_texts)
        return combined_text
    except Exception as e:
        raise ValueError(f"Failed to read Word document (.docx): {str(e)}")


def clean_extracted_text(text: str) -> str:
    """
    Clean and normalize extracted resume text by removing redundant whitespace.

    Args:
        text (str): Raw text string.

    Returns:
        str: Cleaned text string.
    """
    if not text:
        return ""
    
    # Replace null characters or non-printable controls
    text = text.replace("\x00", "")
    # Normalize multiple newlines to max 2
    text = re.sub(r"\n\s*\n", "\n\n", text)
    # Replace multiple spaces with a single space
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def extract_resume_text(uploaded_file) -> str:
    """
    Automatically detect file format (PDF/DOCX) and extract resume text.

    Args:
        uploaded_file: Streamlit UploadedFile object.

    Returns:
        str: Extracted and cleaned text string.

    Raises:
        ValueError: If extraction fails or format is unsupported.
    """
    if uploaded_file is None:
        raise ValueError("No file provided for extraction.")

    file_name = uploaded_file.name
    file_ext = os.path.splitext(file_name)[1].lower()

    # Read bytes from stream safely
    uploaded_file.seek(0)
    file_bytes = uploaded_file.read()

    if file_ext == ".pdf":
        raw_text = extract_text_from_pdf(file_bytes)
    elif file_ext == ".docx":
        raw_text = extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file extension '{file_ext}'.")

    cleaned_text = clean_extracted_text(raw_text)
    return cleaned_text
